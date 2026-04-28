import fitz
from docx import Document

from ddm_engine.extraction.docx_text import DocxTextExtractor
from ddm_engine.extraction.pdf_text import PdfTextExtractor
from ddm_engine.extraction.plain_text import PlainTextExtractor


def test_pdf_text_extractor_returns_page_tokens_with_coordinates() -> None:
    pdf_bytes = _make_pdf_bytes("Jane Doe\njane@example.com")

    layout = PdfTextExtractor().extract("job-1", pdf_bytes)

    assert layout.job_id == "job-1"
    assert layout.source_file_type == "pdf"
    assert len(layout.pages) == 1
    assert layout.token_count >= 2

    page = layout.pages[0]
    assert page.page_number == 1
    assert page.width > 0
    assert page.height > 0

    token_text = {token.text for token in page.tokens}
    assert {"Jane", "Doe"}.issubset(token_text)
    assert all(token.bbox.x1 > token.bbox.x0 for token in page.tokens)
    assert all(token.bbox.y1 > token.bbox.y0 for token in page.tokens)


def test_plain_text_extractor_builds_synthetic_layout() -> None:
    layout = PlainTextExtractor().extract("job-1", b"Jane Doe\njane@example.com")

    assert layout.source_file_type == "txt"
    assert layout.token_count >= 3
    assert layout.pages[0].tokens[0].source == "synthetic_text"


def test_docx_extractor_reads_paragraph_text(tmp_path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com")
    document.save(path)

    layout = DocxTextExtractor().extract("job-1", path.read_bytes())

    assert layout.source_file_type == "docx"
    assert "Jane" in {token.text for token in layout.pages[0].tokens}


def _make_pdf_bytes(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    try:
        return document.tobytes()
    finally:
        document.close()
