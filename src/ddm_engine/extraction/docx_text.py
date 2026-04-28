from __future__ import annotations

from io import BytesIO

from docx import Document

from ddm_engine.extraction.models import DocumentLayout
from ddm_engine.extraction.text_layout import TextLayoutBuilder


class DocxTextExtractor:
    def __init__(self, layout_builder: TextLayoutBuilder | None = None) -> None:
        self.layout_builder = layout_builder or TextLayoutBuilder()

    def extract(self, job_id: str, document_bytes: bytes) -> DocumentLayout:
        document = Document(BytesIO(document_bytes))
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return self.layout_builder.build(job_id, "docx", "\n".join(parts))
